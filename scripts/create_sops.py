"""
Generate the 4 synthetic Biocon SOP .docx files with pre-scripted compliance gaps.

Usage:
    python scripts/create_sops.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


# ---------------------------------------------------------------------------
# BC-MFG-UC-047 — Upstream Cell Culture mAb
# GAP 1: §7.3 no comparability protocol
# GAP 2: §6.2 glycosylation not listed as CQA
# ---------------------------------------------------------------------------

def create_bc_mfg_uc_047(sops_dir: Path) -> None:
    doc = Document()
    doc.add_heading("Standard Operating Procedure", 0)
    para(doc, "Document Number: BC-MFG-UC-047")
    para(doc, "Title: Upstream Cell Culture for Monoclonal Antibody Production")
    para(doc, "Version: 3.2    Effective Date: 01-Mar-2025    Review Date: 01-Mar-2026")
    para(doc, "Department: Manufacturing — Biologics    Site: Biocon Biologics, Bangalore")

    heading(doc, "1. Purpose", 1)
    para(doc,
        "This SOP establishes the procedures for upstream cell culture operations "
        "for monoclonal antibody (mAb) production, including seed train, bioreactor "
        "inoculation, in-process monitoring, and harvest. The purpose is to ensure "
        "consistent, reproducible production of mAb drug substance that meets "
        "pre-defined quality specifications."
    )

    heading(doc, "2. Scope", 1)
    para(doc,
        "This procedure applies to all manufacturing personnel engaged in upstream "
        "cell culture operations for mAb products at the Biocon Biologics Bangalore "
        "facility. It covers operations from cell thaw through bioreactor harvest."
    )

    heading(doc, "3. Definitions", 1)
    para(doc, "3.1 Seed Train: Sequential expansion of cells from frozen vial to production bioreactor.")
    para(doc, "3.2 Cell Viability: Percentage of live cells determined by trypan blue exclusion.")
    para(doc, "3.3 Titre: Concentration of mAb product in harvest material (g/L).")
    para(doc, "3.4 IVCC: Integrated Viable Cell Count, expressed as 10⁶ viable cells·day/mL.")
    para(doc, "3.5 CAPA: Corrective and Preventive Action.")
    para(doc, "3.6 CIP: Clean-in-Place. SIP: Steam-in-Place.")

    heading(doc, "4. Responsibilities", 1)
    para(doc, "4.1 Manufacturing Operator: Performs all cell culture operations per this SOP.")
    para(doc, "4.2 Manufacturing Supervisor: Reviews in-process data and authorises bioreactor inoculation.")
    para(doc, "4.3 Quality Control (QC): Performs and reports all in-process and harvest testing.")
    para(doc, "4.4 Quality Assurance (QA): Reviews batch records, approves deviations, and releases product.")

    heading(doc, "5. Cell Thaw and Seed Train", 1)
    para(doc,
        "5.1 Retrieve the appropriate Working Cell Bank (WCB) vial from cryogenic storage "
        "per SOP BC-MFG-CB-001."
    )
    para(doc,
        "5.2 Thaw vial rapidly at 37°C in a water bath. Transfer contents to a pre-warmed "
        "T-flask containing complete culture medium (CDM4MAb supplemented with 4 mM L-glutamine)."
    )
    para(doc,
        "5.3 Expand cells through a defined seed train: T-flask → spinner flask → wave bag → "
        "N-1 seed bioreactor. Passage at 3–4 × 10⁶ viable cells/mL. Document all passages "
        "in Batch Record BC-MFG-UC-047-BR."
    )
    para(doc,
        "5.4 Inoculate the production bioreactor (N-stage) when the N-1 seed culture reaches "
        "viability ≥ 95% and cell density 4–6 × 10⁶ viable cells/mL."
    )

    heading(doc, "6. Bioreactor Operating Parameters", 1)

    heading(doc, "6.1 In-Process Control Parameters", 2)
    para(doc,
        "The following parameters shall be maintained throughout the production bioreactor run. "
        "Deviations outside these ranges shall trigger a manufacturing deviation per "
        "SOP BC-QA-DEV-001."
    )
    para(doc,
        "• Temperature: 37.0 ± 0.5 °C\n"
        "• pH: 7.00 ± 0.10 (controlled via CO₂ sparging and NaHCO₃ addition)\n"
        "• Dissolved Oxygen (DO): 40 ± 10% air saturation\n"
        "• Agitation: 150–200 RPM\n"
        "• Daily in-process sampling for cell viability, cell density (VCD), lactate, glucose, and osmolality."
    )

    heading(doc, "6.2 Acceptance Criteria for Harvest", 2)
    # GAP 2: glycosylation NOT mentioned as CQA
    para(doc,
        "Harvest shall be initiated when the following criteria are met. Batches not meeting "
        "all acceptance criteria shall be rejected and a deviation report raised."
    )
    para(doc,
        "• Cell viability: ≥ 70% at time of harvest\n"
        "• Product titre: ≥ 2.0 g/L by Protein A HPLC\n"
        "• pH at harvest: 7.0 ± 0.3\n"
        "• Osmolality: 280–350 mOsm/kg\n"
        "• Glucose: ≥ 0.5 g/L\n"
        "• Run duration: ≤ 14 days"
    )
    para(doc,
        "Samples shall be forwarded to QC for bioburden testing and residual host cell protein "
        "(HCP) by ELISA. Results shall be recorded in the Batch Record before harvest proceeds."
    )

    heading(doc, "7. Process Change Control", 1)

    heading(doc, "7.1 Change Initiator Responsibilities", 2)
    para(doc,
        "Any proposed change to the upstream cell culture process must be initiated using "
        "Form BC-QA-CC-001 (Change Control Initiation Form). The initiating department shall "
        "provide a written description of the proposed change, the technical justification, "
        "and the potential impact on product quality."
    )

    heading(doc, "7.2 Impact Assessment Criteria", 2)
    para(doc,
        "The Change Control team shall assess the proposed change against the following criteria:\n"
        "• Impact on cell line stability or growth kinetics\n"
        "• Impact on product titre\n"
        "• Impact on downstream processing compatibility\n"
        "• Regulatory classification (minor, moderate, major) per site regulatory strategy\n"
        "Changes classified as moderate or major shall require QA and Regulatory Affairs review "
        "prior to implementation."
    )

    heading(doc, "7.3 Post-Approval Manufacturing Changes", 2)
    # GAP 1: no comparability protocol, no reference to ICH Q5E or EMA CHMP/437/04
    para(doc,
        "Post-approval manufacturing changes that may affect product quality shall be assessed "
        "for potential impact on safety and efficacy. The change shall be categorised as "
        "minor, moderate, or major in accordance with the applicable regulatory framework. "
        "All changes shall be submitted to Regulatory Affairs for review and, where required, "
        "submitted to the relevant regulatory authority prior to implementation."
    )
    para(doc,
        "The change control record shall document: the nature of the change, the regulatory "
        "classification, the outcome of the internal review, and the date of implementation. "
        "Regulatory Affairs shall maintain a log of all post-approval manufacturing changes."
    )

    heading(doc, "8. Documentation and Records", 1)
    para(doc,
        "8.1 All manufacturing activities shall be recorded contemporaneously in Batch Record "
        "BC-MFG-UC-047-BR. Corrections shall be made by single line through, with date and "
        "initials of the person making the correction.\n"
        "8.2 Completed batch records shall be submitted to QA within 2 business days of harvest.\n"
        "8.3 Electronic data from bioreactor control systems shall be archived per "
        "SOP BC-IT-DS-005."
    )

    heading(doc, "9. Deviations and CAPA", 1)
    para(doc,
        "Any deviation from this SOP shall be reported on Form BC-QA-DEV-001 within 24 hours "
        "of occurrence. Deviations impacting product quality shall trigger a CAPA investigation "
        "per SOP BC-QA-CAPA-001. The Manufacturing Supervisor shall perform an initial impact "
        "assessment and notify QA immediately for deviations classified as critical."
    )

    heading(doc, "10. References", 1)
    para(doc, "• SOP BC-MFG-CB-001: Working Cell Bank Management")
    para(doc, "• SOP BC-QA-DEV-001: Deviation Reporting and Management")
    para(doc, "• SOP BC-QA-CAPA-001: Corrective and Preventive Action")
    para(doc, "• SOP BC-IT-DS-005: Electronic Data Archiving")
    para(doc, "• ICH Q10: Pharmaceutical Quality System")
    para(doc, "• 21 CFR Part 211: Current Good Manufacturing Practice")

    path = sops_dir / "BC-MFG-UC-047_Upstream_Cell_Culture_mAb.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------
# BC-QC-BR-012 — Batch Record Review
# GAP 3: §5.2 review timeline not specified
# ---------------------------------------------------------------------------

def create_bc_qc_br_012(sops_dir: Path) -> None:
    doc = Document()
    doc.add_heading("Standard Operating Procedure", 0)
    para(doc, "Document Number: BC-QC-BR-012")
    para(doc, "Title: Batch Record Review and Approval")
    para(doc, "Version: 2.1    Effective Date: 15-Jan-2025    Review Date: 15-Jan-2026")
    para(doc, "Department: Quality Control / Quality Assurance    Site: Biocon Biologics, Bangalore")

    heading(doc, "1. Purpose", 1)
    para(doc,
        "This SOP defines the process for the review, reconciliation, and approval of "
        "pharmaceutical batch records for biological drug substance and drug product "
        "manufactured at Biocon Biologics. Batch record review is a critical quality "
        "activity that ensures all manufacturing steps have been performed in accordance "
        "with approved procedures and that the batch meets all specifications before release."
    )

    heading(doc, "2. Scope", 1)
    para(doc,
        "This procedure applies to all batch records generated for mAb drug substance "
        "and drug product at the Bangalore facility, including upstream cell culture, "
        "purification, formulation, and fill/finish operations."
    )

    heading(doc, "3. Definitions", 1)
    para(doc, "3.1 Batch Record: The complete record of all manufacturing steps for a specific batch.")
    para(doc, "3.2 Master Batch Record (MBR): The approved template from which executed batch records are produced.")
    para(doc, "3.3 OOS: Out-of-Specification result — a test result outside the approved acceptance criteria.")
    para(doc, "3.4 OOT: Out-of-Trend result — a result within specification but showing an unexpected trend.")
    para(doc, "3.5 Disposition: The formal decision to release, reject, or quarantine a batch.")

    heading(doc, "4. Batch Record Components", 1)
    para(doc,
        "4.1 A complete batch record shall include:\n"
        "• Executed manufacturing instructions with all steps signed and dated\n"
        "• In-process control (IPC) data and results\n"
        "• Environmental monitoring records\n"
        "• Equipment use and cleaning logs\n"
        "• Raw material usage records and Certificate of Analysis\n"
        "• Analytical testing results from QC\n"
        "• Deviation forms (if applicable)\n"
        "• CAPA references (if applicable)"
    )
    para(doc,
        "4.2 Each page of the batch record shall be signed and dated by the operator "
        "performing the step and countersigned by the manufacturing supervisor."
    )

    heading(doc, "5. Review Responsibilities", 1)

    heading(doc, "5.1 QC Analyst Review", 2)
    para(doc,
        "The QC Analyst responsible for testing shall review all analytical results "
        "for accuracy and completeness. The analyst shall:\n"
        "• Verify that all test methods referenced are current and approved versions\n"
        "• Confirm that all results are within approved acceptance criteria or that an "
        "OOS/OOT investigation has been initiated where applicable\n"
        "• Sign and date the QC section of the batch record upon completion of testing\n"
        "• Flag any anomalous results to the QC Supervisor immediately"
    )

    heading(doc, "5.2 Batch Record Review and Approval", 2)
    # GAP 3: no timeline, no deadline, no frequency
    para(doc,
        "Batch records shall be reviewed by the Quality Assurance department prior to "
        "product release. The QA review shall verify that:\n"
        "• All manufacturing steps have been completed as specified in the Master Batch Record\n"
        "• All in-process control results are within approved limits\n"
        "• All deviations have been documented and dispositioned\n"
        "• All required QC testing has been completed and results meet specifications\n"
        "• All required signatures and dates are present and legible"
    )
    para(doc,
        "The Authorised Person (AP) or designee shall sign the batch disposition form "
        "upon satisfactory completion of the QA review, authorising the batch for release "
        "or rejection."
    )

    heading(doc, "5.3 Discrepancy Handling", 2)
    para(doc,
        "Any discrepancy identified during batch record review shall be documented on "
        "Form BC-QA-BRD-001. Minor discrepancies (e.g., missing initials, illegible entries) "
        "shall be resolved by the responsible operator under QA supervision. Major discrepancies "
        "shall trigger a deviation investigation per SOP BC-QA-DEV-001."
    )

    heading(doc, "6. Out-of-Specification Investigations", 1)

    heading(doc, "6.1 OOS Initiation", 2)
    para(doc,
        "Upon receipt of an OOS result, the QC Analyst shall immediately notify the QC "
        "Supervisor and initiate an OOS investigation per SOP BC-QC-OOS-003. The initial "
        "laboratory investigation shall assess whether the OOS result is attributable to a "
        "laboratory error (Phase I investigation) or a genuine product failure (Phase II investigation)."
    )

    heading(doc, "6.2 OOS Documentation", 2)
    para(doc,
        "All OOS investigations shall be fully documented, including the original result, "
        "the investigation findings, retesting results (if applicable), and the final disposition "
        "decision. The batch shall remain on hold until the OOS investigation is completed "
        "and QA has provided written disposition."
    )

    heading(doc, "7. Archiving", 1)
    para(doc,
        "Completed and approved batch records shall be archived in the QA document management "
        "system per SOP BC-QA-DMS-008. Records shall be retained for a minimum of one year "
        "after the expiry date of the batch, or as required by applicable regulations, "
        "whichever is longer."
    )

    heading(doc, "8. References", 1)
    para(doc, "• SOP BC-QA-DEV-001: Deviation Reporting")
    para(doc, "• SOP BC-QC-OOS-003: Out-of-Specification Investigation")
    para(doc, "• SOP BC-QA-DMS-008: Document Management and Archiving")
    para(doc, "• 21 CFR Part 211: Current Good Manufacturing Practice")
    para(doc, "• ICH Q10: Pharmaceutical Quality System")

    path = sops_dir / "BC-QC-BR-012_Batch_Record_Review.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------
# BC-RA-IM-008 — Immunogenicity Risk Assessment
# GAP: §5.1 no tiered testing plan; §5.3 no ADA assay validation parameters
# ---------------------------------------------------------------------------

def create_bc_ra_im_008(sops_dir: Path) -> None:
    doc = Document()
    doc.add_heading("Standard Operating Procedure", 0)
    para(doc, "Document Number: BC-RA-IM-008")
    para(doc, "Title: Immunogenicity Risk Assessment for Biological Products")
    para(doc, "Version: 1.4    Effective Date: 01-Jun-2024    Review Date: 01-Jun-2025")
    para(doc, "Department: Regulatory Affairs    Site: Biocon Biologics, Bangalore")

    heading(doc, "1. Purpose", 1)
    para(doc,
        "This SOP defines the framework for conducting immunogenicity risk assessments "
        "for biological drug products, including monoclonal antibodies (mAbs), developed "
        "or manufactured by Biocon Biologics. The purpose is to identify, assess, and "
        "mitigate the risk of unwanted immune responses in patients, in accordance with "
        "applicable regulatory guidelines."
    )

    heading(doc, "2. Scope", 1)
    para(doc,
        "This procedure applies to all biological product development programmes, "
        "biosimilar development activities, and post-approval immunogenicity monitoring "
        "programmes at Biocon Biologics. It covers pre-clinical, clinical, and "
        "post-marketing immunogenicity assessment activities."
    )

    heading(doc, "3. Regulatory Framework", 1)
    para(doc,
        "Immunogenicity assessments shall be conducted in accordance with applicable "
        "regulatory guidelines, including but not limited to:\n"
        "• EMA Guideline on Immunogenicity Assessment of Therapeutic Proteins (EMEA/CHMP/BMWP/14327/2006)\n"
        "• EMA Guideline on Immunogenicity Assessment of Monoclonal Antibodies (EMA/CHMP/BMWP/86289/2010)\n"
        "• FDA Guidance on Immunogenicity Testing of Therapeutic Protein Products\n"
        "• ICH S6(R1): Preclinical Safety Evaluation of Biotechnology-Derived Pharmaceuticals"
    )

    heading(doc, "4. Risk Factors for Immunogenicity", 1)
    para(doc,
        "4.1 Product-related factors:\n"
        "• Protein sequence and structure (degree of humanisation for mAbs)\n"
        "• Post-translational modifications (glycosylation, aggregation)\n"
        "• Formulation components and excipients\n"
        "• Container closure system and leachables\n"
        "• Degradation products"
    )
    para(doc,
        "4.2 Patient-related factors:\n"
        "• Underlying disease state and immune status\n"
        "• Concomitant immunosuppressive therapy\n"
        "• Genetic factors (HLA haplotype, FcγR polymorphisms)\n"
        "• Prior exposure to related biologics"
    )
    para(doc,
        "4.3 Treatment-related factors:\n"
        "• Route of administration (subcutaneous carries higher immunogenicity risk than IV)\n"
        "• Dose and dosing frequency\n"
        "• Treatment duration and interruptions"
    )

    heading(doc, "5. Immunogenicity Testing Programme", 1)

    heading(doc, "5.1 Pre-clinical Testing Plan", 2)
    # GAP: no tiered testing strategy, no assay format, no cut-point method
    para(doc,
        "Pre-clinical immunogenicity assessment will be conducted as part of the non-clinical "
        "development programme. Immunogenicity will be assessed using industry-standard methods "
        "appropriate for the species and product type. Results will be reported in the relevant "
        "non-clinical study reports and summarised in the regulatory dossier."
    )

    heading(doc, "5.2 Clinical Immunogenicity Monitoring", 2)
    para(doc,
        "5.2.1 Clinical immunogenicity monitoring shall be conducted in all clinical trials "
        "involving the biological product. Serum samples shall be collected at defined timepoints "
        "as specified in the clinical study protocol.\n\n"
        "5.2.2 The immunogenicity testing strategy shall include screening, confirmatory, and "
        "titration assays as applicable. The presence of anti-drug antibodies (ADAs) shall be "
        "assessed in relation to pharmacokinetic data, efficacy endpoints, and safety findings.\n\n"
        "5.2.3 Prevalence and incidence of ADAs shall be reported in clinical study reports "
        "and regulatory submissions in accordance with EMA and FDA guidance."
    )

    heading(doc, "5.3 ADA Assay Validation Requirements", 2)
    # GAP: no drug tolerance, no sensitivity threshold, no precision requirements
    para(doc,
        "All assays used for immunogenicity assessment shall be validated prior to use in "
        "clinical or regulatory studies. Validation shall be performed in accordance with "
        "applicable regulatory guidelines and documented in an assay validation report. "
        "The validated assay shall be used for all clinical sample analyses."
    )

    heading(doc, "6. Risk Mitigation", 1)
    para(doc,
        "6.1 Where immunogenicity risk factors are identified, appropriate risk mitigation "
        "strategies shall be implemented, which may include:\n"
        "• Optimisation of formulation to minimise aggregation\n"
        "• Selection of optimal route of administration\n"
        "• Monitoring protocols for early detection of ADA responses\n"
        "• Risk communication in product labelling"
    )

    heading(doc, "7. Reporting", 1)
    para(doc,
        "7.1 Immunogenicity risk assessments shall be documented in a formal risk assessment "
        "report and reviewed by the Regulatory Affairs team.\n"
        "7.2 Immunogenicity data from clinical trials shall be reported in clinical study "
        "reports and in applicable regulatory submissions (IND, IMPD, BLA, MAA).\n"
        "7.3 Post-approval immunogenicity findings shall be reported per applicable "
        "pharmacovigilance requirements."
    )

    heading(doc, "8. References", 1)
    para(doc, "• EMA BMWP/14327 — Immunogenicity Assessment of Therapeutic Proteins")
    para(doc, "• EMA BMWP/86289 — Immunogenicity Assessment of Monoclonal Antibodies")
    para(doc, "• SOP BC-QC-IM-006: ADA Assay Procedures")
    para(doc, "• SOP BC-RA-DS-002: Regulatory Dossier Management")

    path = sops_dir / "BC-RA-IM-008_Immunogenicity_Risk_Assessment.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------
# BC-AN-MV-031 — Analytical Method Validation — Protein A HPLC
# GAP 4: §6.3 system suitability criteria not defined
# ---------------------------------------------------------------------------

def create_bc_an_mv_031(sops_dir: Path) -> None:
    doc = Document()
    doc.add_heading("Standard Operating Procedure", 0)
    para(doc, "Document Number: BC-AN-MV-031")
    para(doc, "Title: Analytical Method Validation — Residual Protein A Quantification by HPLC")
    para(doc, "Version: 2.0    Effective Date: 01-Sep-2024    Review Date: 01-Sep-2025")
    para(doc, "Department: Analytical Sciences    Site: Biocon Biologics, Bangalore")

    heading(doc, "1. Purpose", 1)
    para(doc,
        "This SOP describes the validation requirements for the HPLC method used to quantify "
        "residual Protein A in mAb drug substance. Protein A is a ligand used in Protein A "
        "affinity chromatography purification, and its residual presence in the final product "
        "must be controlled below defined limits due to its immunogenic potential."
    )

    heading(doc, "2. Scope", 1)
    para(doc,
        "This document applies to the validated HPLC method for residual Protein A quantification "
        "at Biocon Biologics Analytical Sciences laboratories. It covers initial method validation, "
        "revalidation, and method transfer activities."
    )

    heading(doc, "3. Method Description — Protein A HPLC", 1)
    para(doc,
        "3.1 Principle: The method uses reversed-phase HPLC to separate and quantify residual "
        "Protein A in mAb drug substance samples. Protein A is detected by UV absorbance at "
        "280 nm and quantified against a reference standard calibration curve.\n\n"
        "3.2 Column: C18 reversed-phase column (150 mm × 4.6 mm, 3.5 μm particle size).\n\n"
        "3.3 Mobile Phase A: 0.1% trifluoroacetic acid (TFA) in water.\n"
        "Mobile Phase B: 0.1% TFA in acetonitrile.\n\n"
        "3.4 Gradient: 5% B to 65% B over 20 minutes, at a flow rate of 1.0 mL/min.\n\n"
        "3.5 Injection volume: 100 μL. Column temperature: 40°C. Detection: UV 280 nm."
    )

    heading(doc, "4. Validation Parameters", 1)

    heading(doc, "4.1 Specificity", 2)
    para(doc,
        "Specificity shall be demonstrated by showing that the Protein A peak is resolved from "
        "the mAb matrix peak and all known excipients. Resolution (Rs) between the Protein A "
        "peak and the nearest interfering peak shall be ≥ 1.5 under the validated method conditions."
    )

    heading(doc, "4.2 Linearity and Range", 2)
    para(doc,
        "Linearity shall be demonstrated over the validated range of 2 ng/mL to 500 ng/mL "
        "(equivalent to 0.04 ppm to 10 ppm in a 50 g/L drug substance sample). A minimum of "
        "five concentration levels shall be used. The correlation coefficient (r²) shall be ≥ 0.999."
    )

    heading(doc, "4.3 Accuracy and Recovery", 2)
    para(doc,
        "Accuracy shall be assessed by spiking Protein A standard into blank mAb drug substance "
        "matrix at three concentration levels (25%, 100%, and 150% of the target concentration). "
        "A minimum of three replicates at each level shall be prepared. Mean recovery shall be "
        "within 90–110% at each level."
    )

    heading(doc, "4.4 Precision", 2)
    para(doc,
        "4.4.1 Repeatability: Six replicate injections of the same sample preparation at 100% "
        "of the target concentration. %RSD of peak area shall be ≤ 2.0%.\n\n"
        "4.4.2 Intermediate Precision: Three analysts on three different days. The overall "
        "%RSD of results across analysts and days shall be ≤ 5.0%."
    )

    heading(doc, "4.5 Limit of Detection and Limit of Quantitation", 2)
    para(doc,
        "LOD shall be determined as the concentration giving a signal-to-noise ratio of 3:1. "
        "LOQ shall be determined as the concentration giving a signal-to-noise ratio of 10:1 "
        "and a %RSD ≤ 10% across six replicate injections. The validated LOQ shall be ≤ 2 ng/mL."
    )

    heading(doc, "5. Equipment and Reagents", 1)
    para(doc,
        "5.1 HPLC system with UV detector, autosampler, and column oven (Agilent 1260 Infinity "
        "or equivalent).\n"
        "5.2 Analytical balance, calibrated to ± 0.1 mg.\n"
        "5.3 Protein A reference standard (certificate of analysis required, purity ≥ 95%).\n"
        "5.4 HPLC-grade acetonitrile, TFA, and water.\n"
        "5.5 Volumetric flasks and pipettes, calibrated."
    )

    heading(doc, "6. Procedure", 1)

    heading(doc, "6.1 Sample Preparation", 2)
    para(doc,
        "6.1.1 Dilute drug substance to 50 g/L with mobile phase A.\n"
        "6.1.2 Centrifuge at 10,000 × g for 5 minutes if the sample is turbid.\n"
        "6.1.3 Transfer 100 μL to an HPLC vial. Analyse within 24 hours of preparation."
    )

    heading(doc, "6.2 Calibration Standards", 2)
    para(doc,
        "6.2.1 Prepare a stock solution of Protein A reference standard at 10 μg/mL in "
        "mobile phase A.\n"
        "6.2.2 Prepare calibration standards at 2, 10, 50, 100, 250, and 500 ng/mL by "
        "serial dilution of the stock solution.\n"
        "6.2.3 Run calibration standards at the beginning and end of each analytical sequence. "
        "The calibration curve must meet linearity criteria (r² ≥ 0.999) for results to be "
        "considered valid."
    )

    heading(doc, "6.3 System Suitability", 2)
    # GAP 4: system suitability mentioned but no criteria defined
    para(doc,
        "System suitability shall be confirmed prior to sample analysis. The system suitability "
        "test shall be performed by injecting the system suitability standard at the beginning "
        "of each analytical sequence. System suitability results shall be recorded in the "
        "analytical worksheet and reviewed by the QC analyst before proceeding with sample "
        "analysis. If system suitability is not met, the cause shall be investigated before "
        "sample analysis begins."
    )

    heading(doc, "7. Calculations", 1)
    para(doc,
        "7.1 Protein A concentration in the sample shall be calculated from the calibration "
        "curve using linear regression.\n"
        "7.2 Results shall be expressed as ng/mL in the diluted sample and converted to ppm "
        "(ng/mg) relative to drug substance concentration.\n"
        "7.3 The acceptance limit for residual Protein A in drug substance is ≤ 10 ppm."
    )

    heading(doc, "8. Method Transfer", 1)
    para(doc,
        "8.1 Prior to transfer, the sending laboratory shall provide the receiving laboratory "
        "with the validated method documentation, reference standards, and historical "
        "performance data.\n"
        "8.2 Method transfer shall be conducted per SOP BC-AN-MT-002."
    )

    heading(doc, "9. References", 1)
    para(doc, "• SOP BC-AN-MT-002: Analytical Method Transfer Procedure")
    para(doc, "• ICH Q2(R1): Validation of Analytical Procedures — Text and Methodology")
    para(doc, "• 21 CFR 211.194: Laboratory Records")
    para(doc, "• USP <621>: Chromatography")

    path = sops_dir / "BC-AN-MV-031_Analytical_Method_Validation_ProteinA_HPLC.docx"
    doc.save(str(path))
    print(f"  Created {path.name}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    sops_dir = Path(__file__).resolve().parent.parent / "sops"
    sops_dir.mkdir(exist_ok=True)

    print("Creating synthetic SOP files in sops/ …")
    create_bc_mfg_uc_047(sops_dir)
    create_bc_qc_br_012(sops_dir)
    create_bc_ra_im_008(sops_dir)
    create_bc_an_mv_031(sops_dir)
    print(f"\nDone — 4 SOP files written to {sops_dir}")


if __name__ == "__main__":
    main()
