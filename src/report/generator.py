"""
Gap Report Generator — reads output/gap_registry.json and renders an
audit-ready DOCX report. Never calls any external API.

Usage:
    python -m src.report.generator --registry output/gap_registry.json
    python -m src.report.generator --registry output/gap_registry.json --output output/my_report.docx
"""
from __future__ import annotations

import argparse
import json
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from src.schemas import GapRegistry, GapResult, Severity


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

COLOUR_CRITICAL = RGBColor(0xC0, 0x00, 0x00)   # dark red
COLOUR_MAJOR = RGBColor(0xED, 0x7D, 0x31)       # orange
COLOUR_MINOR = RGBColor(0x70, 0xAD, 0x47)       # green
COLOUR_HEADING = RGBColor(0x1F, 0x49, 0x7D)     # dark blue

SEVERITY_COLOURS = {
    Severity.CRITICAL: COLOUR_CRITICAL,
    Severity.MAJOR: COLOUR_MAJOR,
    Severity.MINOR: COLOUR_MINOR,
}


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _add_horizontal_line(doc: Document) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell: Any, hex_colour: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _bold_para(doc: Document, text: str, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_value = p.add_run(value)
    run_value.font.size = Pt(9)


def _blockquote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text[:500] + ("…" if len(text) > 500 else ""))
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


# ---------------------------------------------------------------------------
# Report generator
# ---------------------------------------------------------------------------

class ReportGenerator:
    def __init__(self, registry_path: Path) -> None:
        self.registry_path = registry_path
        data = json.loads(registry_path.read_text())
        self.registry = GapRegistry.model_validate(data)

    def render(self, output_path: Path) -> Path:
        doc = Document()
        self._set_margins(doc)
        self._cover_page(doc)
        doc.add_page_break()
        self._executive_summary(doc)
        doc.add_page_break()
        self._checklist_summary(doc)
        doc.add_page_break()
        self._gap_registry_section(doc)
        self._appendix_methodology(doc)

        output_path.parent.mkdir(exist_ok=True)
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _cover_page(self, doc: Document) -> None:
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_heading("SOP Compliance Gap Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Pharmaceutical Regulatory Compliance Assessment")
        run.font.size = Pt(14)
        run.font.color.rgb = COLOUR_HEADING

        doc.add_paragraph()
        _add_horizontal_line(doc)
        doc.add_paragraph()

        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = "Table Grid"
        meta_data = [
            ("Prepared for", "Biocon Biologics Limited"),
            ("Prepared by", "EMB Global"),
            ("Report date", date.today().strftime("%d-%b-%Y")),
            ("SOPs assessed", str(self.registry.total_sops_scanned)),
            ("Total gaps identified", str(self.registry.total_gaps_found)),
        ]
        for row, (label, value) in zip(meta_table.rows, meta_data):
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
            _set_cell_bg(row.cells[0], "EBF3FB")

        doc.add_paragraph()
        _add_horizontal_line(doc)
        doc.add_paragraph()

        classif = doc.add_paragraph()
        classif.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = classif.add_run("CONFIDENTIAL — FOR INTERNAL USE ONLY")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOUR_CRITICAL

    # ------------------------------------------------------------------
    # Executive summary
    # ------------------------------------------------------------------

    def _executive_summary(self, doc: Document) -> None:
        doc.add_heading("Executive Summary", level=1)

        # Count severities across all SOPs
        total_critical = total_major = total_minor = 0
        for scan in self.registry.scans:
            for f in scan.findings:
                if f.severity == Severity.CRITICAL:
                    total_critical += 1
                elif f.severity == Severity.MAJOR:
                    total_major += 1
                else:
                    total_minor += 1

        p = doc.add_paragraph()
        p.add_run(
            f"This report presents the findings of an automated regulatory compliance gap analysis "
            f"conducted against {self.registry.total_sops_scanned} Biocon Biologics Standard Operating "
            f"Procedures. The analysis identified a total of {self.registry.total_gaps_found} compliance "
            f"gap(s): "
        )
        r_crit = p.add_run(f"{total_critical} CRITICAL")
        r_crit.bold = True
        r_crit.font.color.rgb = COLOUR_CRITICAL
        p.add_run(", ")
        r_maj = p.add_run(f"{total_major} MAJOR")
        r_maj.bold = True
        r_maj.font.color.rgb = COLOUR_MAJOR
        p.add_run(", and ")
        r_min = p.add_run(f"{total_minor} MINOR")
        r_min.bold = True
        r_min.font.color.rgb = COLOUR_MINOR
        p.add_run(".")

        doc.add_paragraph()

        # Summary table
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["SOP ID", "SOP File", "Clauses Scanned", "CRITICAL", "MAJOR", "MINOR"]
        header_row = table.rows[0]
        for i, hdr in enumerate(headers):
            header_row.cells[i].text = hdr
            header_row.cells[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(header_row.cells[i], "1F497D")
            header_row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for scan in sorted(self.registry.scans, key=lambda s: s.sop_id):
            crit = sum(1 for f in scan.findings if f.severity == Severity.CRITICAL)
            maj = sum(1 for f in scan.findings if f.severity == Severity.MAJOR)
            minor = sum(1 for f in scan.findings if f.severity == Severity.MINOR)
            row = table.add_row()
            row.cells[0].text = scan.sop_id
            row.cells[1].text = scan.sop_file
            row.cells[2].text = str(scan.total_clauses_scanned)
            row.cells[3].text = str(crit)
            row.cells[4].text = str(maj)
            row.cells[5].text = str(minor)
            if crit > 0:
                _set_cell_bg(row.cells[3], "FFCCCC")
            if maj > 0:
                _set_cell_bg(row.cells[4], "FFE4CC")

    # ------------------------------------------------------------------
    # Compliance checklist summary
    # ------------------------------------------------------------------

    def _checklist_summary(self, doc: Document) -> None:
        doc.add_heading("Compliance Checklist Summary", level=1)
        intro = doc.add_paragraph(
            "Each row represents one SOP clause reviewed. "
            "A clause is marked PASS when no gap was identified against the regulatory corpus. "
            "Any clause with one or more findings is marked with its highest severity."
        )
        intro.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        PASS_BG   = "EBF5EB"
        CRIT_BG   = "FFDDDD"
        MAJOR_BG  = "FFEEDD"
        MINOR_BG  = "FAFFF0"
        HDR_BG    = "002F59"

        for scan in sorted(self.registry.scans, key=lambda s: s.sop_id):
            doc.add_heading(f"{scan.sop_id}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"File: {scan.sop_file}   |   "
                      f"Clauses reviewed: {scan.total_clauses_scanned}   |   "
                      f"Gaps found: {scan.gaps_found}").font.size = Pt(8)

            # Build clause → worst finding map
            clause_map: dict[str, GapResult] = {}
            for f in scan.findings:
                existing = clause_map.get(f.sop_clause)
                order = {"CRITICAL": 0, "MAJOR": 1, "MINOR": 2}
                if existing is None or order[f.severity.value] < order[existing.severity.value]:
                    clause_map[f.sop_clause] = f

            # Collect all clause ids that appeared (from findings; gaps-only list)
            # Supplemented with PASS rows for clauses without findings
            all_gap_clauses = sorted(clause_map.keys())

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0]
            for i, txt in enumerate(["Clause", "Description / Gap", "Regulation", "Status"]):
                hdr.cells[i].text = txt
                run = hdr.cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(hdr.cells[i], HDR_BG)

            # Rows with gaps
            for clause_id in all_gap_clauses:
                finding = clause_map[clause_id]
                sev = finding.severity.value
                row = table.add_row()

                row.cells[0].text = clause_id
                row.cells[1].text = finding.gap_description[:180] + (
                    "…" if len(finding.gap_description) > 180 else ""
                )
                row.cells[2].text = finding.regulation_ref
                row.cells[3].text = sev

                bg = {"CRITICAL": CRIT_BG, "MAJOR": MAJOR_BG, "MINOR": MINOR_BG}.get(sev, "FFFFFF")
                colour = SEVERITY_COLOURS[finding.severity]
                for cell in row.cells:
                    _set_cell_bg(cell, bg)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                # Bold + colour the status cell
                status_run = row.cells[3].paragraphs[0].runs[0]
                status_run.bold = True
                status_run.font.color.rgb = colour

            # PASS row summary
            pass_count = scan.total_clauses_scanned - len(all_gap_clauses)
            if pass_count > 0:
                row = table.add_row()
                row.cells[0].text = f"({pass_count} clause(s))"
                row.cells[1].text = "No regulatory gaps identified"
                row.cells[2].text = "—"
                row.cells[3].text = "PASS"
                for cell in row.cells:
                    _set_cell_bg(cell, PASS_BG)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                pass_run = row.cells[3].paragraphs[0].runs[0]
                pass_run.bold = True
                pass_run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

            doc.add_paragraph()

    # ------------------------------------------------------------------
    # Per-SOP gap sections
    # ------------------------------------------------------------------

    def _gap_registry_section(self, doc: Document) -> None:
        doc.add_heading("Gap Registry", level=1)

        for scan in sorted(self.registry.scans, key=lambda s: s.sop_id):
            doc.add_heading(f"{scan.sop_id} — {scan.sop_file}", level=2)
            p = doc.add_paragraph()
            p.add_run(
                f"Scan timestamp: {scan.scan_timestamp}  |  "
                f"Clauses scanned: {scan.total_clauses_scanned}  |  "
                f"Gaps found: {scan.gaps_found}"
            ).font.size = Pt(8)

            if not scan.findings:
                doc.add_paragraph("No compliance gaps detected for this SOP.")
                continue

            for i, finding in enumerate(
                sorted(scan.findings, key=lambda f: (f.severity.value, f.sop_clause))
            ):
                self._render_finding(doc, i + 1, finding)
                _add_horizontal_line(doc)

    def _render_finding(self, doc: Document, n: int, finding: GapResult) -> None:
        colour = SEVERITY_COLOURS[finding.severity]

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run_num = p.add_run(f"Finding #{n}  ")
        run_num.bold = True
        run_sev = p.add_run(f"[{finding.severity.value}]")
        run_sev.bold = True
        run_sev.font.color.rgb = colour
        p.add_run(f"  {finding.sop_clause} — {finding.regulation_ref}")

        _label_value(doc, "Clause", finding.sop_clause)
        _label_value(doc, "Regulation", finding.regulation_ref)
        _label_value(doc, "Confidence", f"{finding.confidence:.0%}")

        _label_value(doc, "Gap Description", finding.gap_description)
        _label_value(doc, "Remediation", finding.remediation)

        p_sop = doc.add_paragraph()
        p_sop.paragraph_format.space_before = Pt(4)
        run = p_sop.add_run("SOP Clause Text (verbatim):")
        run.bold = True
        run.font.size = Pt(9)
        _blockquote(doc, finding.sop_clause_text)

        p_reg = doc.add_paragraph()
        run = p_reg.add_run("Regulatory Requirement (verbatim):")
        run.bold = True
        run.font.size = Pt(9)
        _blockquote(doc, finding.regulation_excerpt)

    # ------------------------------------------------------------------
    # Appendix — methodology
    # ------------------------------------------------------------------

    def _appendix_methodology(self, doc: Document) -> None:
        doc.add_page_break()
        doc.add_heading("Appendix — Methodology", level=1)

        doc.add_heading("Pipeline Overview", level=2)
        doc.add_paragraph(
            "This report was generated by the Biocon SOP Compliance Engine, a RAG-based "
            "automated regulatory compliance analysis system. The pipeline operates as follows:"
        )
        for step in [
            "1. Regulatory PDFs are loaded, chunked (512 tokens / 64-token overlap), and embedded "
            "   using OpenAI text-embedding-3-large into a local ChromaDB vector store.",
            "2. Each SOP is parsed into individual clauses by section number.",
            "3. For each clause, the top-8 most relevant regulatory chunks are retrieved by "
            "   cosine similarity search.",
            "4. Claude (via OpenRouter) analyses each clause against the retrieved regulatory "
            "   context at temperature=0 and returns structured gap findings.",
            "5. Findings are validated against the GapResult schema and aggregated into this report.",
        ]:
            doc.add_paragraph(step, style="List Bullet")

        doc.add_heading("Model and Parameters", level=2)
        doc.add_paragraph(
            f"LLM: anthropic/claude-sonnet-4-5 via OpenRouter\n"
            f"Temperature: 0 (deterministic output)\n"
            f"Embedding model: text-embedding-3-large\n"
            f"Chunk size: 512 tokens  |  Overlap: 64 tokens\n"
            f"Retrieval top-k: 8  |  Minimum relevance score: 0.30"
        )

        doc.add_heading("Severity Definitions", level=2)
        for sev, defn in [
            ("CRITICAL", "Missing required element; will cause regulatory non-conformance at inspection."),
            ("MAJOR", "Deficient element; likely to be cited at FDA/EMA inspection."),
            ("MINOR", "Improvement recommended; low regulatory risk; not typically cited."),
        ]:
            p = doc.add_paragraph()
            run = p.add_run(f"{sev}: ")
            run.bold = True
            run.font.color.rgb = SEVERITY_COLOURS[Severity(sev)]
            p.add_run(defn)

    # ------------------------------------------------------------------
    # Page margins
    # ------------------------------------------------------------------

    def _set_margins(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate audit-ready DOCX gap report from gap_registry.json"
    )
    parser.add_argument(
        "--registry",
        type=Path,
        default=Path("output/gap_registry.json"),
        help="Path to gap_registry.json (default: output/gap_registry.json)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Output .docx path (default: output/gap_report_YYYY-MM-DD.docx)",
    )
    args = parser.parse_args()

    if not args.registry.exists():
        raise FileNotFoundError(
            f"Registry not found: {args.registry}\n"
            "Run the detector first: python -m src.gap_engine.detector --all --sops sops/"
        )

    output_path = args.output or (
        Path("output") / f"gap_report_{date.today().isoformat()}.docx"
    )

    gen = ReportGenerator(args.registry)
    out = gen.render(output_path)
    print(f"Report written to: {out}")


if __name__ == "__main__":
    main()
